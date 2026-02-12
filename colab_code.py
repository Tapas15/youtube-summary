# ============================================================
# YouTube to Book Summary Converter - Colab Code
# ============================================================

# CELL 1: Install Dependencies
!pip install -q youtube-transcript-api groq python-docx reportlab

# CELL 2: Configure API Key
import os
GROQ_API_KEY = ""  # <-- ENTER YOUR API KEY HERE
os.environ["GROQ_API_KEY"] = GROQ_API_KEY
print("API Key configured!" if GROQ_API_KEY else "Please enter your API Key above!")

# CELL 3: Enter YouTube URL
YOUTUBE_URL = ""  # <-- ENTER YOUR YOUTUBE URL HERE
print(f"Ready: {YOUTUBE_URL}" if YOUTUBE_URL else "Please enter a YouTube URL above!")

# CELL 4: Import Libraries
import re
from datetime import datetime, timedelta
from youtube_transcript_api import YouTubeTranscriptApi
from groq import Groq
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER
from IPython.display import display, Markdown
print("Libraries imported!")

# CELL 5: Statistics Calculator
def calculate_transcript_stats(text, video_duration_seconds=None):
    """Calculate comprehensive statistics for a transcript."""
    if not text:
        return {
            'word_count': 0,
            'character_count': 0,
            'reading_time_minutes': 0,
            'reading_time_formatted': '0 seconds',
            'video_duration_seconds': None,
            'video_duration_formatted': None,
            'speaking_rate_wpm': None,
            'paragraph_count': 0,
            'sentence_count': 0,
        }
    
    # Basic counts
    word_count = len([w for w in text.split() if w])
    char_count = len(text)
    paragraph_count = max(1, len([p for p in text.split('\n\n') if p.strip()]))
    sentence_count = max(1, len([s for s in re.split(r'[.!?]+', text) if s.strip()]))
    
    # Reading time (at 200 WPM)
    reading_time_minutes = round(word_count / 200, 1)
    if reading_time_minutes < 1:
        reading_time_formatted = f"{int(reading_time_minutes * 60)} seconds"
    elif reading_time_minutes < 60:
        reading_time_formatted = f"{reading_time_minutes:.1f} minutes"
    else:
        hours = int(reading_time_minutes // 60)
        mins = int(reading_time_minutes % 60)
        reading_time_formatted = f"{hours}h {mins}m"
    
    # Video duration formatting
    video_duration_formatted = None
    if video_duration_seconds:
        td = timedelta(seconds=int(video_duration_seconds))
        total_seconds = int(td.total_seconds())
        hours, remainder = divmod(total_seconds, 3600)
        minutes, secs = divmod(remainder, 60)
        if hours > 0:
            video_duration_formatted = f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            video_duration_formatted = f"{minutes:02d}:{secs:02d}"
    
    # Speaking rate
    speaking_rate = None
    if video_duration_seconds and video_duration_seconds > 0:
        speaking_rate = round(word_count / (video_duration_seconds / 60))
    
    return {
        'word_count': word_count,
        'character_count': char_count,
        'reading_time_minutes': reading_time_minutes,
        'reading_time_formatted': reading_time_formatted,
        'video_duration_seconds': video_duration_seconds,
        'video_duration_formatted': video_duration_formatted,
        'speaking_rate_wpm': speaking_rate,
        'paragraph_count': paragraph_count,
        'sentence_count': sentence_count,
    }

print("Statistics calculator ready!")

# CELL 6: Extract Transcript
def get_video_id(url):
    patterns = [
        r'(?:v=|/)([0-9A-Za-z_-]{11}).*',
        r'(?:youtu\.be/)([0-9A-Za-z_-]{11})',
        r'(?:youtube\.com/shorts/)([0-9A-Za-z_-]{11})',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def extract_transcript(url):
    video_id = get_video_id(url)
    if not video_id:
        return None, "Invalid URL"
    try:
        api = YouTubeTranscriptApi()
        tl = api.list(video_id)
        try:
            t = tl.find_manually_created_transcript(['en'])
        except:
            t = tl.find_generated_transcript(['en'])
        data = t.fetch()
        
        # Extract text and calculate duration
        text_parts = []
        total_duration = 0.0
        for e in data:
            text = e.text if hasattr(e, 'text') else str(e)
            # Clean up timestamp markers
            if '[' in text and ']' in text:
                text = re.sub(r'\[.*?\]', '', text).strip()
            text_parts.append(text)
            
            # Calculate video duration
            if hasattr(e, 'start') and hasattr(e, 'duration'):
                end_time = e.start + e.duration
                if end_time > total_duration:
                    total_duration = end_time
        
        text = ' '.join(text_parts)
        return text, {'video_id': video_id, 'duration': total_duration}
    except Exception as e:
        return None, str(e)

print("Functions ready!")

# CELL 7: Generate Summary
def create_summary(text, api_key):
    client = Groq(api_key=api_key)
    prompt = f"""You are an expert author. Create a comprehensive book-style summary with:

1. EXECUTIVE OVERVIEW (2-3 paragraphs)
2. INTRODUCTION
3. CHAPTER-BY-CHAPTER SUMMARY
4. KEY CONCEPTS
5. KEY TAKEAWAYS (numbered)
6. MEMORABLE QUOTATIONS
7. PRACTICAL APPLICATIONS
8. CRITICAL ANALYSIS
9. CONCLUSION

Write in flowing prose.

TRANSCRIPT:
{text[:15000]}"""
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": "You are an expert author"},
            {"role": "user", "content": prompt}
        ],
        max_tokens=8000,
        temperature=0.5
    )
    return response.choices[0].message.content

print("Summary function ready!")

# CELL 8: Save All Files (Word + PDF)
def save_all_files(summary, transcript, video_id, title, stats):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = "".join(c for c in title if c.isalnum() or c in ' -_').strip()[:40]
    files_created = []
    
    # === 1. Save Book as Word ===
    book_file = f"{safe_title}_book_{timestamp}.docx"
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Video ID: {video_id}")
    doc.add_paragraph(f"Generated: {datetime.now().isoformat()}")
    doc.add_paragraph()
    
    sections = summary.split('\n# ')
    for section in sections:
        if section.strip():
            lines = section.strip().split('\n')
            section_title = lines[0].replace('#', '').strip()
            if section_title:
                doc.add_heading(section_title, level=1)
            content = '\n'.join(lines[1:]).strip()
            if content:
                doc.add_paragraph(content)
    doc.save(book_file)
    files_created.append(("📕 Book (Word)", book_file))
    
    # === 2. Save Book as PDF ===
    pdf_file = f"{safe_title}_book_{timestamp}.pdf"
    doc_pdf = SimpleDocTemplate(pdf_file, pagesize=letter,
                                leftMargin=0.5*inch, rightMargin=0.5*inch,
                                topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='CenterTitle', parent=styles['Title'], alignment=TA_CENTER))
    styles.add(ParagraphStyle(name='SectionHeader', parent=styles['Heading1'], fontSize=14))
    styles.add(ParagraphStyle(name='BodyText', parent=styles['Normal'], fontSize=11, leading=14))
    
    story = []
    story.append(Paragraph(title, styles['CenterTitle']))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Video ID: {video_id}", styles['BodyText']))
    story.append(Paragraph(f"Generated: {datetime.now().isoformat()}", styles['BodyText']))
    story.append(Spacer(1, 24))
    
    for section in sections:
        if section.strip():
            lines = section.strip().split('\n')
            section_title = lines[0].replace('#', '').strip()
            if section_title:
                story.append(Paragraph(section_title, styles['SectionHeader']))
            content = '\n'.join(lines[1:]).strip()
            if content:
                for para in content.split('\n\n'):
                    if para.strip():
                        story.append(Paragraph(para, styles['BodyText']))
            story.append(Spacer(1, 12))
    
    doc_pdf.build(story)
    files_created.append(("📗 Book (PDF)", pdf_file))
    
    # === 3. Save Transcript as Word ===
    trans_file = f"{safe_title}_transcript_{timestamp}.docx"
    doc2 = Document()
    
    # Title
    title_para = doc2.add_heading(f"Transcript: {title}", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Video Metadata Section
    doc2.add_paragraph()
    doc2.add_heading("VIDEO METADATA", level=1)
    doc2.add_paragraph(f"Video ID: {video_id}")
    if stats.get('video_duration_formatted'):
        doc2.add_paragraph(f"Video Duration: {stats['video_duration_formatted']}")
    doc2.add_paragraph(f"Generated: {datetime.now().isoformat()}")
    
    # Transcript Statistics Section
    doc2.add_paragraph()
    doc2.add_heading("TRANSCRIPT STATISTICS", level=1)
    doc2.add_paragraph(f"Word Count: {stats['word_count']:,} words")
    doc2.add_paragraph(f"Character Count: {stats['character_count']:,} characters")
    doc2.add_paragraph(f"Reading Time: {stats['reading_time_formatted']} (at 200 WPM)")
    if stats.get('speaking_rate_wpm'):
        doc2.add_paragraph(f"Speaking Rate: {stats['speaking_rate_wpm']} words/min")
    doc2.add_paragraph(f"Paragraphs: {stats['paragraph_count']}")
    doc2.add_paragraph(f"Sentences: {stats['sentence_count']}")
    
    doc2.add_paragraph()
    doc2.add_paragraph("─" * 50)
    doc2.add_paragraph()
    
    # Transcript content
    max_chars = 15000
    chunks = [transcript[i:i+max_chars] for i in range(0, len(transcript), max_chars)]
    
    for i, chunk in enumerate(chunks):
        if i > 0:
            doc2.add_paragraph()
            doc2.add_run().add_break(WD_BREAK.PAGE)
            doc2.add_paragraph()
        doc2.add_heading(f"Page {i+1} of {len(chunks)}", level=2)
        doc2.add_paragraph(chunk)
    
    # Final Summary Page
    doc2.add_paragraph()
    doc2.add_run().add_break(WD_BREAK.PAGE)
    doc2.add_heading("TRANSCRIPT SUMMARY", level=1)
    doc2.add_paragraph()
    doc2.add_paragraph(f"Total Words: {stats['word_count']:,}")
    doc2.add_paragraph(f"Total Characters: {stats['character_count']:,}")
    doc2.add_paragraph(f"Estimated Read Time: {stats['reading_time_formatted']} (at 200 WPM)")
    if stats.get('video_duration_formatted'):
        doc2.add_paragraph(f"Video Duration: {stats['video_duration_formatted']}")
    if stats.get('speaking_rate_wpm'):
        doc2.add_paragraph(f"Speaking Pace: {stats['speaking_rate_wpm']} words/min")
    doc2.add_paragraph(f"Pages Generated: {len(chunks)}")
    
    doc2.save(trans_file)
    files_created.append(("📘 Transcript (Word)", trans_file))
    
    return files_created

print("File saving functions ready!")

# CELL 9: RUN EVERYTHING
from google.colab import files

if YOUTUBE_URL and GROQ_API_KEY:
    print("Processing... Please wait...\n")
    
    # Extract
    transcript, result = extract_transcript(YOUTUBE_URL)
    
    if isinstance(result, str) and transcript is None:
        print(f"Error: {result}")
    else:
        video_id = result['video_id']
        video_duration = result.get('duration', 0)
        print(f"Transcript: {len(transcript)} chars")
        print(f"Video Duration: {video_duration:.1f} seconds")
        
        # Calculate statistics
        stats = calculate_transcript_stats(transcript, video_duration)
        print(f"\n📊 Transcript Statistics:")
        print(f"   Word Count: {stats['word_count']:,}")
        print(f"   Reading Time: {stats['reading_time_formatted']}")
        if stats['speaking_rate_wpm']:
            print(f"   Speaking Rate: {stats['speaking_rate_wpm']} WPM")
        
        # Summary
        print("\n📚 Generating summary...")
        summary = create_summary(transcript, GROQ_API_KEY)
        print("Summary generated!")
        
        # Save files
        title = f"YouTube Video {video_id}"
        files_created = save_all_files(summary, transcript, video_id, title, stats)
        
        print(f"\n📁 Files saved:")
        for label, filename in files_created:
            print(f"   {label}: {filename}")
        
        print("\n" + "="*60)
        print("BOOK SUMMARY")
        print("="*60)
        display(Markdown("# " + title + "\n\n" + summary))
        
        print("\n📥 Download files:")
        for label, filename in files_created:
            files.download(filename)

else:
    print("Enter API Key and YouTube URL above!")

# ============================================================
