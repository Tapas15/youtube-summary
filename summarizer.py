"""Main summarizer module that orchestrates the entire process."""
import os
from typing import Optional, Dict, Any, Tuple
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER

from config import config
from youtube_transcript import YouTubeExtractor, extract_transcript
from groq_client import GroqClient, create_client
from utils.statistics import calculate_transcript_stats


class YouTubeBookSummarizer:
    """Main class for converting YouTube videos to book-style summaries."""
    
    def __init__(
        self,
        api_key: Optional[str] = None,
        model: Optional[str] = None,
        output_dir: Optional[Path] = None,
    ):
        """
        Initialize the summarizer.
        
        Args:
            api_key: Groq API key
            model: LLM model to use
            output_dir: Directory for saving summaries
        """
        self.transcript_extractor = YouTubeExtractor()
        self.groq_client = create_client(api_key=api_key, model=model)
        self.output_dir = output_dir or config.OUTPUT_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Load the summary template
        self.template = self._load_template()
    
    def _load_template(self) -> str:
        """Load the summary template from youtube_summary.md."""
        template_path = Path(__file__).parent / "youtube_summary.md"
        if template_path.exists():
            return template_path.read_text()
        return self._get_default_template()
    
    def _get_default_template(self) -> str:
        """Get default template if file not found."""
        return """# IDENTITY and PURPOSE

You are an expert author and content analyst. You transform YouTube video transcripts into comprehensive, well-structured book-style summaries that capture the full depth and nuance of the content.

# STEPS

1. READ AND COMPREHEND THE ENTIRE TRANSCRIPT:
   - First read through the entire transcript to understand the overall structure
   - Identify the main topic, key themes, and primary arguments
   - Note the speaker's expertise and perspective
   - Identify any supporting evidence, examples, or anecdotes

2. ANALYZE THE CONTENT STRUCTURE:
   - Identify major sections or chapters in the video
   - Extract key turning points or major topic shifts
   - Note recurring themes or concepts
   - Identify the core message the presenter is conveying

3. EXTRACT AND ORGANIZE KEY INFORMATION:
   - Pull out main ideas and central arguments
   - Extract important facts, statistics, and data points
   - Note expert opinions or expert quotes mentioned
   - Identify practical applications or takeaways
   - Capture stories, examples, and case studies

4. SYNTHESIZE INTO BOOK FORMAT:
   - Organize content into proper book structure
   - Write in narrative prose, avoiding excessive lists
   - Maintain the presenter's voice and perspective
   - Connect ideas logically across sections

5. WRITE COMPREHENSIVE SUMMARY WITH:

# OUTPUT SECTIONS

- EXECUTIVE OVERVIEW:
  - What this video/book covers in 2-3 paragraphs
  - Why this content matters
  - Who would benefit most from this

- INTRODUCTION:
  - Background and context of the topic
  - Presenter's credentials and perspective
  - What problem or question the content addresses

- CHAPTER-BY-CHAPTER SUMMARY:
  **Chapter 1: [Chapter Title]**
  - Detailed summary of this section
  - Key concepts introduced
  - Main arguments or points made
  - Examples or evidence provided
  - How this chapter builds on previous content

  **Chapter 2: [Chapter Title]**
  - [Same structure]

- KEY CONCEPTS AND DEFINITIONS:
  - Important terms explained in context
  - Technical vocabulary with plain-language explanations
  - Conceptual frameworks introduced

- KEY TAKEAWAYS (numbered):
  1. [Most important insight]
  2. [Second most important insight]
  3. [Third most important insight]

- MEMORABLE QUOTATIONS:
  - "Quote from the video that captures essential wisdom"
  - "Another powerful statement"

- PRACTICAL APPLICATIONS:
  - How to apply the concepts learned
  - Actionable steps or exercises
  - Real-world scenarios where this knowledge applies

- CRITICAL ANALYSIS:
  - Strengths of the arguments presented
  - Potential limitations or counterpoints
  - Areas where more depth would be valuable

- FURTHER READING/SOURCES:
  - Books or resources mentioned in the video
  - Related topics for exploration

- CONCLUSION:
  - Synthesis of the overall message
  - Final wisdom or call to action
  - How this content fits into larger context

- QUICK BULLET SUMMARY (TL;DR):
  - 5-10 bullet points capturing the absolute most important takeaways
  - Each bullet should be a complete, standalone insight
  - Written for quick scanning and easy understanding
  - Focus on actionable insights and key learnings
  - Example format:
    • [Main point 1 - the core message in one sentence]
    • [Main point 2 - key insight or finding]
    • [Main point 3 - practical takeaway]
    • [Main point 4 - important statistic or fact]
    • [Main point 5 - recommended action or next step]

# OUTPUT INSTRUCTIONS

- WRITE IN PROSE: Use flowing paragraphs, not excessive bullet points
- MAINTAIN DEPTH: Each chapter summary should be 3-5 substantial paragraphs
- USE CONSISTENT FORMATTING: Follow the structure above exactly
- INCLUDE SPECIFICS: Include actual examples, numbers, and quotes from the content
- BOOK-LIKE QUALITY: Write as if creating a valuable reference book
- MAXIMIZE VALUE: Extract every useful insight from the transcript
- DO NOT REPEAT: Each piece of information should appear only once
- VARIED OPENINGS: Don't start multiple sections with the same words
- COMPLETE THOUGHTS: Each paragraph should fully develop an idea
- NO PLACEHOLDERS: Fill in actual titles, quotes, and content from the transcript
"""
    
    def summarize(
        self,
        url: str,
        language: str = "en",
        save_output: bool = True,
        verbose: bool = True,
    ) -> Dict[str, Any]:
        """
        Generate a book-style summary from a YouTube video.
        
        Args:
            url: YouTube video URL
            language: Transcript language
            save_output: Whether to save summary to file
            verbose: Print progress messages
            
        Returns:
            Dictionary with summary and metadata
        """
        result = {
            'success': False,
            'video_url': url,
            'summary': None,
            'metadata': {},
            'error': None
        }
        
        try:
            # Step 1: Extract transcript
            if verbose:
                print("[Step 1] Extracting transcript from YouTube...")
            
            transcript_result = extract_transcript(url, language)
            
            if not transcript_result['success']:
                raise ValueError(transcript_result['error'])
            
            result['metadata']['transcript'] = transcript_result['metadata']
            transcript = transcript_result['text']
            
            if verbose:
                print(f"  [OK] Transcript extracted: {transcript_result['metadata']['transcript_length']} characters")
                print(f"  [OK] Language: {transcript_result['metadata']['language']}")
            
            # Step 2: Process transcript with Groq
            if verbose:
                print("[Step 2] Generating book-style summary with AI...")
            
            # Check if we need to chunk the transcript
            if len(transcript) > config.MAX_TRANSCRIPT_LENGTH:
                if verbose:
                    print(f"  [INFO] Transcript is long ({len(transcript)} chars), processing in chunks...")
                
                chunk_summaries = self._process_in_chunks(transcript, verbose)
                llm_result = self.groq_client.summarize_transcript(
                    transcript=transcript,
                    prompt_template=self.template,
                    chunk_summaries=chunk_summaries,
                )
            else:
                llm_result = self.groq_client.summarize_transcript(
                    transcript=transcript,
                    prompt_template=self.template,
                )
            
            if not llm_result['success']:
                raise ValueError(llm_result['error'])
            
            result['summary'] = llm_result['content']
            result['metadata']['usage'] = llm_result['usage']
            result['metadata']['model'] = llm_result['model']
            result['success'] = True
            
            if verbose:
                print(f"  [OK] Summary generated!")
                print(f"  [OK] Tokens used: {llm_result['usage']['total_tokens']}")
                print(f"  [OK] Estimated cost: ${llm_result['usage']['estimated_cost_usd']:.4f}")
            
            # Step 3: Save output
            if save_output:
                if verbose:
                    print("[Step 3] Saving files...")
                
                # Save both book summary and transcript
                output_paths = self._save_book_summary(
                    result['summary'],
                    transcript_result['text'],
                    transcript_result['metadata'],
                )
                # output_paths is a tuple: (book_path, pdf_path, transcript_path)
                result['metadata']['output_file'] = output_paths
                
                if verbose:
                    book_path, pdf_path, transcript_path = output_paths
                    print(f"  [OK] Files saved!")
                    print(f"    - Book (Word): {book_path}")
                    print(f"    - Book (PDF): {pdf_path}")
                    print(f"    - Transcript: {transcript_path}")
            
            if verbose:
                print("\n[SUCCESS] Summary complete!")
            
            return result
            
        except Exception as e:
            result['error'] = str(e)
            if verbose:
                print(f"\n[ERROR] Error: {str(e)}")
            return result
    
    def _process_in_chunks(
        self,
        transcript: str,
        verbose: bool = True,
    ) -> list:
        """
        Process long transcripts in chunks.
        
        Args:
            transcript: Full transcript text
            verbose: Print progress
            
        Returns:
            List of chunk summaries
        """
        extractor = YouTubeExtractor()
        chunks = extractor.chunk_transcript(
            text=transcript,
            max_length=config.MAX_TRANSCRIPT_LENGTH,
            overlap=config.CHUNK_OVERLAP
        )
        
        if verbose:
            print(f"  [INFO] Split into {len(chunks)} chunks")
        
        chunk_summaries = []
        
        for i, chunk in enumerate(chunks):
            if verbose:
                print(f"     Processing chunk {i + 1}/{len(chunks)}...")
            
            # Create a focused summary for each chunk
            chunk_prompt = f"""Summarize the following transcript segment. Identify:
1. Main topics discussed
2. Key points and arguments
3. Important examples or evidence
4. Any memorable quotes

Return this as a concise but comprehensive summary.

TRANSCRIPT SEGMENT:
{chunk}
"""
            
            result = self.groq_client.chat_completion(
                messages=[{"role": "user", "content": chunk_prompt}],
                max_tokens=4000,
                temperature=0.5,
            )
            
            if result['success']:
                chunk_summaries.append(result['content'])
            else:
                # Fall back to raw chunk if API fails
                chunk_summaries.append(f"[Segment {i + 1}] {chunk[:500]}...")
        
        return chunk_summaries
    
    def _save_book_summary(
        self,
        summary: str,
        transcript: str,
        metadata: Dict[str, Any],
    ) -> Tuple[Path, Path, Path]:
        """Save both book summary and transcript to Word files.
        
        Returns:
            Tuple of (book_path, pdf_path, transcript_path)
        """
        video_id = metadata.get('video_id', 'unknown')
        title = metadata.get('title', f'YouTube Video {video_id}')
        
        # Create filename base
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = "".join(
            c for c in title
            if c.isalnum() or c in (' ', '-', '_')
        ).strip()[:50]
        
        # Save Book Summary as Word file
        book_filename = f"{safe_title}_{video_id}_book_{timestamp}.docx"
        book_path = self.output_dir / book_filename
        
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Video ID: {video_id}")
        doc.add_paragraph(f"Generated: {datetime.now().isoformat()}")
        doc.add_paragraph(f"Transcript Length: {metadata.get('transcript_length', 'N/A')} characters")
        doc.add_paragraph()
        
        # Add summary sections
        sections = summary.split('\n# ')
        for section in sections:
            if section.strip():
                lines = section.strip().split('\n')
                if lines:
                    # First line is the section title
                    section_title = lines[0].replace('#', '').strip()
                    if section_title:
                        doc.add_heading(section_title, level=1)
                    
                    # Rest is content
                    content = '\n'.join(lines[1:]).strip()
                    if content:
                        doc.add_paragraph(content)
        
        doc.save(book_path)
        
        # Also save as PDF
        pdf_filename = f"{safe_title}_{video_id}_book_{timestamp}.pdf"
        pdf_path = self.output_dir / pdf_filename
        
        # Create PDF
        doc_pdf = SimpleDocTemplate(str(pdf_path), pagesize=letter,
                                    leftMargin=0.5*inch, rightMargin=0.5*inch,
                                    topMargin=0.5*inch, bottomMargin=0.5*inch)
        
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='CenterTitle', parent=styles['Title'], alignment=TA_CENTER))
        styles.add(ParagraphStyle(name='SectionHeader', parent=styles['Heading1'], fontSize=14, spaceAfter=12))
        # Use existing BodyText style or create custom one
        if 'CustomBody' not in styles.byName:
            styles.add(ParagraphStyle(name='CustomBody', parent=styles['Normal'], fontSize=11, leading=14))
        
        story = []
        
        # Title
        story.append(Paragraph(title, styles['CenterTitle']))
        story.append(Spacer(1, 12))
        
        # Metadata
        story.append(Paragraph(f"Video ID: {video_id}", styles['CustomBody']))
        story.append(Paragraph(f"Generated: {datetime.now().isoformat()}", styles['CustomBody']))
        story.append(Spacer(1, 24))
        
        # Summary sections
        sections = summary.split('\n# ')
        for section in sections:
            if section.strip():
                lines = section.strip().split('\n')
                if lines:
                    section_title = lines[0].replace('#', '').strip()
                    if section_title:
                        story.append(Paragraph(section_title, styles['SectionHeader']))
                    
                    content = '\n'.join(lines[1:]).strip()
                    if content:
                        # Wrap long text
                        for para in content.split('\n\n'):
                            if para.strip():
                                story.append(Paragraph(para, styles['CustomBody']))
                    
                    story.append(Spacer(1, 12))
        
        doc_pdf.build(story)
        
        # Save Transcript as Word file
        transcript_filename = f"{safe_title}_{video_id}_transcript_{timestamp}.docx"
        transcript_path = self.output_dir / transcript_filename
        
        doc_transcript = Document()
        
        # Add title for transcript
        title_para2 = doc_transcript.add_heading(f"Transcript: {title}", 0)
        title_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Calculate transcript statistics
        video_duration = metadata.get('video_duration_seconds')
        stats = calculate_transcript_stats(transcript, video_duration)
        
        # Add metadata section
        doc_transcript.add_paragraph()
        meta_heading = doc_transcript.add_heading("VIDEO METADATA", level=1)
        
        # Video metadata table
        doc_transcript.add_paragraph(f"Video ID: {video_id}")
        if stats.get('video_duration_formatted'):
            doc_transcript.add_paragraph(f"Video Duration: {stats['video_duration_formatted']}")
        doc_transcript.add_paragraph(f"Language: {metadata.get('language', 'en')}")
        doc_transcript.add_paragraph(f"Generated: {datetime.now().isoformat()}")
        
        # Add statistics section
        doc_transcript.add_paragraph()
        stats_heading = doc_transcript.add_heading("TRANSCRIPT STATISTICS", level=1)
        
        doc_transcript.add_paragraph(f"Word Count: {stats['word_count']:,} words")
        doc_transcript.add_paragraph(f"Character Count: {stats['character_count']:,} characters")
        doc_transcript.add_paragraph(f"Reading Time: {stats['reading_time_formatted']} (at 200 WPM)")
        if stats.get('speaking_rate_wpm'):
            doc_transcript.add_paragraph(f"Speaking Rate: {stats['speaking_rate_wpm']} words/min")
        doc_transcript.add_paragraph(f"Paragraphs: {stats['paragraph_count']}")
        doc_transcript.add_paragraph(f"Sentences: {stats['sentence_count']}")
        
        doc_transcript.add_paragraph()
        doc_transcript.add_paragraph("─" * 50)
        doc_transcript.add_paragraph()
        
        # Add transcript content (split into pages)
        max_chars = 15000  # Approximate chars per page
        chunks = [transcript[i:i+max_chars] for i in range(0, len(transcript), max_chars)]
        
        for i, chunk in enumerate(chunks):
            if i > 0:
                # Add page break using paragraph with break
                para = doc_transcript.add_paragraph()
                para.add_run().add_break(WD_BREAK.PAGE)
            
            # Page header
            page_header = doc_transcript.add_heading(f"Page {i+1} of {len(chunks)}", level=2)
            
            doc_transcript.add_paragraph(chunk)
        
        # Add final summary statistics page
        para = doc_transcript.add_paragraph()
        para.add_run().add_break(WD_BREAK.PAGE)
        
        final_heading = doc_transcript.add_heading("TRANSCRIPT SUMMARY", level=1)
        doc_transcript.add_paragraph()
        doc_transcript.add_paragraph(f"Total Words: {stats['word_count']:,}")
        doc_transcript.add_paragraph(f"Total Characters: {stats['character_count']:,}")
        doc_transcript.add_paragraph(f"Estimated Read Time: {stats['reading_time_formatted']} (at 200 WPM)")
        if stats.get('video_duration_formatted'):
            doc_transcript.add_paragraph(f"Video Duration: {stats['video_duration_formatted']}")
        if stats.get('speaking_rate_wpm'):
            doc_transcript.add_paragraph(f"Speaking Pace: {stats['speaking_rate_wpm']} words/min")
        doc_transcript.add_paragraph(f"Pages Generated: {len(chunks)}")
        
        doc_transcript.save(transcript_path)
        
        # Return all file paths
        return book_path, pdf_path, transcript_path

    def _save_summary(self, summary: str, metadata: Dict[str, Any]) -> Path:
        """Legacy function - kept for compatibility."""
        video_id = metadata.get('video_id', 'unknown')
        title = metadata.get('title', f'YouTube Video {video_id}')
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = "".join(
            c for c in title
            if c.isalnum() or c in (' ', '-', '_')
        ).strip()[:50]
        filename = f"{safe_title}_{video_id}_{timestamp}.md"
        
        output_path = self.output_dir / filename
        
        header = f"""---
title: "{title}"
video_id: "{video_id}"
generated: "{datetime.now().isoformat()}"
transcript_length: "{metadata.get('transcript_length', 'N/A')}"
language: "{metadata.get('language', 'en')}"
---

"""
        
        full_content = header + summary
        output_path.write_text(full_content)
        
        return output_path


def create_summarizer(
    api_key: Optional[str] = None,
    model: Optional[str] = None,
) -> YouTubeBookSummarizer:
    """Factory function to create summarizer instance."""
    return YouTubeBookSummarizer(api_key=api_key, model=model)


if __name__ == "__main__":
    import sys
    
    # Example usage
    summarizer = create_summarizer()
    
    if len(sys.argv) > 1:
        url = sys.argv[1]
        result = summarizer.summarize(url, verbose=True)
        
        if result['success']:
            print(f"\n{'='*60}")
            print("SUMMARY PREVIEW:")
            print(f"{'='*60}")
            print(result['summary'][:1000] + "..." if len(result['summary']) > 1000 else result['summary'])
    else:
        print("Usage: python summarizer.py <YouTube_URL>")
